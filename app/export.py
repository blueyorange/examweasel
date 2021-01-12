import docx
from docx.shared import Cm
from os import getcwd, path
from flask import current_app

def word(questions):
    '''Takes in a list of path names to images to return Word
    Doc with images of questions and mark schemes.
    questions: a list in form [{"q",[path,...]},{"ms",[path,...},...]'''
    print(questions)
    filename = 'New_Document.docx'
    filepath = path.join(current_app.config['DOWNLOAD_FOLDER'],filename)
    # Create word document object
    doc = docx.Document()
    resource_types = {'qp','ms'}
    headings = {'qp' : '' ,'ms' : "Mark Schemes"}
    # Create table with two columns and 'number of questions' rows
    for resource_type in resource_types:
        # Add heading if necessary
        if headings[resource_type]:
            doc.add_heading(headings[resource_type], 0)
        table = doc.add_table(1,2)
        for question in questions:
            image_paths = question[resource_type]
            if image_paths:
                for image_path in image_paths:
                    row_cells = table.add_row().cells
                    if image_paths.index(image_path) == 0:
                        question_number = questions.index(question)+1
                        print( str(question_number) )
                        row_cells[0].text = str(question_number)
                    paragraph = row_cells[1].add_paragraph()
                    run = paragraph.add_run()
                    run.add_picture(image_path,width=Cm(16.5))
    doc.save(filepath)
    return filename
